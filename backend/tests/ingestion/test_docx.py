from pathlib import Path

import docx

from app.ingestion.base import BlockType
from app.ingestion.docx import DocxExtractor


def _build_docx(path: Path) -> None:
    document = docx.Document()
    document.add_paragraph("Q3 2024 Investor Update")
    document.add_paragraph("Revenue grew 12% year over year.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Revenue"
    table.rows[1].cells[1].text = "12345"
    document.save(path)


def test_extracts_paragraphs_and_table(tmp_path: Path) -> None:
    docx_path = tmp_path / "memo.docx"
    _build_docx(docx_path)

    blocks = DocxExtractor().extract(docx_path, document_id="doc-1")

    text_blocks = [b for b in blocks if b.block_type == BlockType.TEXT]
    table_blocks = [b for b in blocks if b.block_type == BlockType.TABLE]

    assert any("Investor Update" in b.text for b in text_blocks)
    assert any("Revenue grew" in b.text for b in text_blocks)
    assert all(b.page == 1 for b in blocks)
    assert all(b.document_id == "doc-1" for b in blocks)

    assert len(table_blocks) == 1
    assert table_blocks[0].table == [["Metric", "Value"], ["Revenue", "12345"]]
