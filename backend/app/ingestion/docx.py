"""DOCX extractor: paragraph text and tables via python-docx. Docx has no
native page concept, so every block is stamped `page=1`."""

from __future__ import annotations

from pathlib import Path

import docx

from app.ingestion.base import BlockType, ContentBlock, Extractor


class DocxExtractor(Extractor):
    def extract(self, file_path: Path, document_id: str) -> list[ContentBlock]:
        document = docx.Document(str(file_path))
        blocks: list[ContentBlock] = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                blocks.append(
                    ContentBlock(
                        document_id=document_id,
                        page=1,
                        block_type=BlockType.TEXT,
                        text=paragraph.text,
                    )
                )

        for table in document.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            if rows:
                blocks.append(
                    ContentBlock(
                        document_id=document_id,
                        page=1,
                        block_type=BlockType.TABLE,
                        table=rows,
                    )
                )

        return blocks
